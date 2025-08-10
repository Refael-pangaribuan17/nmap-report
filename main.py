#!/usr/bin/env python3
import os
import re
import subprocess
from colorama import Fore, Style, init
import pyfiglet
from docx import Document
from fpdf import FPDF

init(autoreset=True)

def print_logo():
    logo = r'''
   .:'                                  `:.
  ::'                                    `::
 :: :.                                  .: ::

  `:. `:.             .             .:'  .:'
   `::. `::           !           ::' .::'
       `::.`::.    .' ! `.    .::'.::'
         `:.  `::::'':!:``::::'   ::'
         :'*:::.  .:' ! `:.  .:::*`:
        :: HHH::.   ` ! '   .::HHH ::
       ::: `H TH::.  `!'  .::HT H' :::
       ::..  `THHH:`:   :':HHHT'  ..::
       `::      `T: `. .' :T'      ::'
         `:. .   :         :   . .:'
           `::'               `::'
             :'  .`.  .  .'.  `:
             :' ::.       .:: `:
             :' `:::     :::' `:
              `.  ``     ''  .'
               :`...........':
               ` :`.     .': '
                `:  `"""'  :'
    '''
    print(Fore.GREEN + logo)
    print(Fore.YELLOW + Style.BRIGHT + pyfiglet.figlet_format("Nmap Report", font="slant"))

target = None
scan_flag = None

def show_nmap_help():
    doc = f'''
{Fore.CYAN}[Nmap - Network Mapper]{Style.RESET_ALL}
Nmap adalah tool scanning jaringan dan port yang sangat populer dan powerful.

Contoh penggunaan dan opsi scan:

  1)  -F                     : Quick Scan, scan 100 port TCP paling umum
  2)  -sV                    : Deteksi versi service yang berjalan pada port terbuka
  3)  -p-                    : Scan semua port TCP (1-65535)
  4)  -A                     : Aggressive Scan (deteksi versi, OS, traceroute, script default)
  5)  --script default       : Jalankan NSE Script default (discovery dan info)
  6)  -O                     : Deteksi sistem operasi target (butuh root)
  7)  -Pn                    : Skip host discovery, anggap host hidup (no ping)
  8)  -sU                    : Scan UDP port
  9)  -sS                    : TCP SYN Scan (stealth)
 10)  -sT                    : TCP Connect Scan (full TCP connect)
 11)  -sX                    : Xmas Scan (stealth, TCP flags FIN, PSH, URG)
 12)  -sN                    : Null Scan (tidak ada flag TCP)
 13)  -sF                    : FIN Scan (flag FIN TCP)
 14)  -sV --script=default   : Scan versi & jalankan script NSE default
 15)  -tr                    : Lakukan traceroute ke target
 16)  -sn                    : Ping Scan, hanya cek host hidup tanpa scan port
 17)  -sO                    : Protocol Scan, cek protokol IP yang didukung target
 18)  --traceroute            : Traceroute (alternatif -tr)
 19)  --top-ports <num>       : Scan top <num> port terbuka paling umum
 20)  -6                     : Scan IPv6 address
 21)  -sI <zombie host>       : Idle Scan (spoof scan, stealth)
 22)  --reason                : Tampilkan alasan status port
 23)  --open                  : Tampilkan hanya port terbuka
 24)  --max-retries <num>     : Batasi jumlah retry paket
 25)  --max-rate <num>        : Batasi rate paket per detik

{Fore.MAGENTA}Catatan:{Style.RESET_ALL} Gunakan Nmap sesuai izin hukum. Scanning tanpa izin dapat melanggar hukum.
'''
    print(doc)

def input_target():
    global target
    target = input(Fore.CYAN + "[?] Masukkan IP/Hostname/URL target: " + Style.RESET_ALL).strip()
    if target == "":
        print(Fore.RED + "[!] Target tidak boleh kosong.")
        target = None
    else:
        print(Fore.GREEN + f"[+] Target disimpan: {target}")

def pilih_scan():
    global scan_flag
    print(Fore.YELLOW + "Pilih jenis scan:")
    print("  1) Quick Scan (-F)")
    print("  2) Service Version (-sV)")
    print("  3) Full Scan (all ports -p-)")
    print("  4) Aggressive Scan (-A)")
    print("  5) NSE Script Scan (--script default)")
    print("  6) OS Detection (-O)")
    print("  7) Skip Host Discovery (-Pn)")
    print("  8) UDP Scan (-sU)")
    print("  9) TCP SYN Scan (-sS)")
    print(" 10) TCP Connect Scan (-sT)")
    print(" 11) Xmas Scan (-sX)")
    print(" 12) Null Scan (-sN)")
    print(" 13) FIN Scan (-sF)")
    print(" 14) Version & Script Scan (-sV --script=default)")
    print(" 15) Traceroute (-tr)")
    print(" 16) Ping Scan (-sn)")
    print(" 17) Protocol Scan (-sO)")
    print(" 18) Traceroute (--traceroute)")
    print(" 19) Top Ports Scan (--top-ports 100)")
    print(" 20) IPv6 Scan (-6)")
    print(" 21) Idle Scan (-sI <zombie host>)")
    print(" 22) Show Reason (--reason)")
    print(" 23) Show Only Open Ports (--open)")
    print(" 24) Max Retries (--max-retries 3)")
    print(" 25) Max Rate (--max-rate 1000)")

    pilihan = input(Fore.CYAN + "Pilih (1-25): " + Style.RESET_ALL)

    if pilihan == "1":
        scan_flag = "-F"
    elif pilihan == "2":
        scan_flag = "-sV"
    elif pilihan == "3":
        scan_flag = "-p- -T4 --min-rate 1000"
    elif pilihan == "4":
        scan_flag = "-A"
    elif pilihan == "5":
        scan_flag = "--script default"
    elif pilihan == "6":
        scan_flag = "-O"
    elif pilihan == "7":
        scan_flag = "-Pn"
    elif pilihan == "8":
        scan_flag = "-sU"
    elif pilihan == "9":
        scan_flag = "-sS"
    elif pilihan == "10":
        scan_flag = "-sT"
    elif pilihan == "11":
        scan_flag = "-sX"
    elif pilihan == "12":
        scan_flag = "-sN"
    elif pilihan == "13":
        scan_flag = "-sF"
    elif pilihan == "14":
        scan_flag = "-sV --script=default"
    elif pilihan == "15":
        scan_flag = "-tr"
    elif pilihan == "16":
        scan_flag = "-sn"
    elif pilihan == "17":
        scan_flag = "-sO"
    elif pilihan == "18":
        scan_flag = "--traceroute"
    elif pilihan == "19":
        scan_flag = "--top-ports 100"
    elif pilihan == "20":
        scan_flag = "-6"
    elif pilihan == "21":
        print(Fore.YELLOW + "Idle scan butuh zombie host (komputer target untuk spoof).")
        zombie = input(Fore.CYAN + "Masukkan IP zombie host: " + Style.RESET_ALL).strip()
        if zombie:
            scan_flag = f"-sI {zombie}"
        else:
            print(Fore.RED + "[!] Zombie host tidak boleh kosong.")
            return
    elif pilihan == "22":
        scan_flag = "--reason"
    elif pilihan == "23":
        scan_flag = "--open"
    elif pilihan == "24":
        scan_flag = "--max-retries 3"
    elif pilihan == "25":
        scan_flag = "--max-rate 1000"
    else:
        print(Fore.RED + "[!] Pilihan tidak valid.")
        return
    print(Fore.GREEN + f"[+] Opsi scan disimpan: {scan_flag}")

def simpan_laporan_docx(data, filename):
    doc = Document()
    doc.add_heading('Nmap Scan Report', 0)
    doc.add_paragraph(data)
    doc.save(filename)

def simpan_laporan_pdf(data, filename):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for line in data.split('\n'):
        pdf.cell(0, 8, txt=line, ln=1)
    pdf.output(filename)

def simpan_laporan(data):
    if not os.path.exists("reports"):
        os.makedirs("reports")
    safe_target = re.sub(r'[^a-zA-Z0-9._-]', '_', target)
    docx_file = f"reports/report_{safe_target}.docx"
    pdf_file = f"reports/report_{safe_target}.pdf"
    simpan_laporan_docx(data, docx_file)
    simpan_laporan_pdf(data, pdf_file)
    print(Fore.GREEN + f"[+] Laporan disimpan di: {docx_file} dan {pdf_file}")

def run_scan():
    global target, scan_flag
    if not target:
        print(Fore.RED + "[!] Target belum diinput.")
        return
    if not scan_flag:
        print(Fore.RED + "[!] Jenis scan belum dipilih.")
        return

    if target.startswith("http://") or target.startswith("https://"):
        print(Fore.RED + "[!] Hapus protokol (http:// atau https://) dari target, hanya domain/IP saja.")
        return

    cmd = f"nmap {scan_flag} {target}"
    print(Fore.YELLOW + f"[+] Running: {cmd} ... (tunggu sebentar)")
    try:
        hasil = subprocess.check_output(cmd, shell=True, text=True)
        simpan_laporan(hasil)
    except subprocess.CalledProcessError as e:
        print(Fore.RED + "[!] Error saat menjalankan scan.")
        if e.output:
            print(e.output)

def main():
    while True:
        os.system("clear")
        print_logo()
        print(Fore.YELLOW + Style.BRIGHT + "[1] Input target (IP/Hostname/URL)")
        print("[2] Lihat penjelasan fitur Nmap")
        print("[3] Pilih jenis scan Nmap")
        print("[4] Jalankan scan dan buat laporan")
        print("[5] Keluar")
        pilihan = input(Fore.CYAN + "Pilih menu (1-5): " + Style.RESET_ALL)

        if pilihan == "1":
            input_target()
        elif pilihan == "2":
            show_nmap_help()
            input(Fore.CYAN + "\nTekan Enter untuk kembali...")
        elif pilihan == "3":
            pilih_scan()
        elif pilihan == "4":
            run_scan()
            input(Fore.CYAN + "\nTekan Enter untuk kembali...")
        elif pilihan == "5":
            print(Fore.YELLOW + "Keluar program. Terima kasih!")
            break
        else:
            print(Fore.RED + "[!] Pilihan tidak valid.")
            input(Fore.CYAN + "\nTekan Enter untuk kembali...")

if __name__ == "__main__":
    main()
